from docx import Document       # package <python-docx>
# from docx.enum.text import WD_BREAK
from docx.shared import Pt
import util

def create_word(law_contents, date_from, date_to):
    doc = Document()
    #스타일 적용하기
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)

    # section = doc.sections[0]
    # header = section.header
    # hdr_p = header.paragraphs[0]
    # hdr_p.text = "시행공포법령"

    # 제목
    doc.add_heading('시행공포법령', level=0)
    doc.add_heading("검색 기간: " + util.format_date(date_from) \
                          + " ~ " + util.format_date(date_to), level=2)
    doc.add_heading("동 기간 중 시행일 법령 및 공포일 법령 검색", level=2)

    for item in law_contents:
        doc.add_page_break()
        doc.add_heading(str(item["법령명_한글"]) + '\n', level=1)
        # doc.add_paragraph("[법령일련번호] " + str(item["법령일련번호"]), style='Intense Quote')
        doc.add_paragraph("[시행 " + util.format_date(str(item["시행일자"])) + "]" \
                          + " [" + str(item["법종구분"]) + " " + util.format_number(str(item["공포번호"])) + ", "
                          + "공포 " + util.format_date(str(item["공포일자"])) \
                          + ", " + str(item["제개정구분"]) + "]", style='Intense Quote')
        p = doc.add_paragraph()
        p.add_run(item["제개정이유내용"])
    
    return doc